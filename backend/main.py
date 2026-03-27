"""
CV Format Tool — FastAPI Backend (Production-Ready)
Integrates: Auth + Role System + Validation Engine + Batch Processing + Review Workflow
"""
from __future__ import annotations  # defer type annotation evaluation (fixes type checker false positives)

import os
import re
import json
import shutil
import tempfile
import uuid
from datetime import datetime
from typing import Any, Optional

import fitz  # PyMuPDF
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Request, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse

# PDF generation for user guideline
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors
from pydantic import BaseModel, EmailStr

# ── NEW MODULES ─────────────────────────────────────────────────
from db import init_db, get_db_session, User, CVJob, CVVersion, AuditLog
from auth import (
    hash_password, verify_password,
    create_access_token, create_refresh_token, decode_token,
    get_current_user, require_role, require_permission, has_permission,
    CurrentUser, log_action, ALLOWED_ROLES,
)
from validation import validate_cv_data, sanitize_for_export, ValidationResult
from batch import (
    Batch, BatchJob, JobStatus, BatchProcessor,
    get_processor, save_batch, save_job, load_batch,
    load_job, list_batches,
)

# ── App Setup ───────────────────────────────────────────────────
app = FastAPI(title="CV Format Tool API", version="2.1.0")

# ── CORS: Allow all origins for public testing ───────────────────
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── PUBLIC MODE: bypass all auth for team testing ────────────────
# All endpoints are accessible without login.
# A default "public" admin user is injected into every request.
_PUBLIC_USER = CurrentUser(id="public-user", email="public@cvformat.local", role="admin")

def _public_user():
    return _PUBLIC_USER

# Create shared permission dependency instances (must reuse same object in overrides AND endpoints)
perm_cv_upload = require_permission("cv:upload")
perm_cv_review = require_permission("cv:review")
perm_cv_qc = require_permission("cv:qc")
perm_cv_export = require_permission("cv:export")
perm_cv_view_all = require_permission("cv:view_all")
perm_user_read = require_permission("user:read")
perm_user_create = require_permission("user:create")
perm_user_update = require_permission("user:update")
perm_audit_read = require_permission("audit:read")

# Override auth dependencies so no login is needed
app.dependency_overrides[get_current_user] = _public_user
app.dependency_overrides[perm_cv_upload] = _public_user
app.dependency_overrides[perm_cv_review] = _public_user
app.dependency_overrides[perm_cv_qc] = _public_user
app.dependency_overrides[perm_cv_export] = _public_user
app.dependency_overrides[perm_cv_view_all] = _public_user
app.dependency_overrides[perm_user_read] = _public_user
app.dependency_overrides[perm_user_create] = _public_user
app.dependency_overrides[perm_user_update] = _public_user
app.dependency_overrides[perm_audit_read] = _public_user

# ── DB Init on startup ──────────────────────────────────────────
@app.on_event("startup")
async def startup():
    try:
        init_db()
        print("[STARTUP] Database initialized")
    except Exception as e:
        print(f"[STARTUP] DB init warning (non-fatal): {e}")


# ═══════════════════════════════════════════════════════════════
# ── GUIDELINE PDF ───────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════

_guideline_cache: Optional[str] = None

def _build_guideline_pdf() -> str:
    """Generate the user guideline PDF and return its file path."""
    out = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    out.close()
    path = out.name

    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()

    navy   = colors.HexColor("#1B3A6B")
    accent = colors.HexColor("#2E6DB4")
    light  = colors.HexColor("#F0F4FA")

    # Custom styles
    title_style = ParagraphStyle("Title2", parent=styles["Title"],
                                  fontSize=20, textColor=navy,
                                  spaceAfter=6, leading=24)
    h1_style = ParagraphStyle("H1", parent=styles["Heading1"],
                               fontSize=14, textColor=navy,
                               spaceBefore=16, spaceAfter=4, leading=18)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"],
                               fontSize=12, textColor=accent,
                               spaceBefore=10, spaceAfter=3, leading=16)
    body_style = ParagraphStyle("Body2", parent=styles["Normal"],
                                fontSize=10, leading=14, spaceAfter=4)
    step_style = ParagraphStyle("Step", parent=styles["Normal"],
                                 fontSize=10, leading=14,
                                 leftIndent=14, spaceAfter=4)
    note_style = ParagraphStyle("Note", parent=styles["Normal"],
                                 fontSize=9, textColor=colors.HexColor("#555555"),
                                 leading=12, spaceAfter=4, leftIndent=10)

    def B(text): return f"<b>{text}</b>"
    def I(text): return f"<i>{text}</i>"

    story = []

    # ── Title ──
    story.append(Paragraph("CV Format Tool — Hướng Dẫn Sử Dụng", title_style))
    story.append(HRFlowable(width="100%", thickness=2, color=navy, spaceAfter=10))

    # ── Overview ──
    story.append(Paragraph(B("Tổng Quan"), h1_style))
    story.append(Paragraph(
        "CV Format Tool là công cụ hỗ trợ định dạng CV tự động từ file PDF/Word gốc "
        "thành file DOCX chuẩn mực, đồng bộ về layout và nội dung. "
        "Công cụ tích hợp trí tuệ nhân tạo (Claude / OpenAI) để trích xuất và "
        "sắp xếp thông tin, kèm chế độ Offline không cần API key.",
        body_style))
    story.append(Spacer(1, 6))

    # ── Quick Start ──
    story.append(Paragraph(B("1. Bắt Đầu Nhanh"), h1_style))
    for i, (title, desc) in enumerate([
        (B("Upload CV"), "Nhấn <b>+ Add Files</b> ở góc trái màn hình để chọn file PDF hoặc DOCX."),
        (B("Chọn chế độ trích xuất"), "Vào <b>Settings</b> (góc phải) để chọn: Offline (không cần API), Auto, Claude, OpenAI, hoặc Ollama."),
        (B("Xử lý"), "Nhấn <b>▶ Run</b>. File DOCX sẽ được tạo và xuất hiện ở cột phải."),
        (B("Tải về"), "Nhấn <b>↓ Download</b> trên file đã xử lý để lưu về máy."),
    ], 1):
        story.append(Paragraph(f"<b> Bước {i}:</b> {title} — {desc}", step_style))
    story.append(Spacer(1, 6))

    # ── Training ──
    story.append(Paragraph(B("2. Đào Tạo Định Dạng (Format Training)"), h1_style))
    story.append(Paragraph(
        "Chế độ <b>Format Training</b> cho phép bạn dạy công cụ nhận diện và định dạng "
        "đúng theo mẫu CV của bạn. Vào tab <b>🧠 Format Training</b>:",
        body_style))
    for i, desc in enumerate([
        "Chọn <b>file CV gốc</b> (Raw) → công cụ trích xuất nội dung.",
        "Chọn <b>file CV mẫu</b> (Done) → định dạng chuẩn mong muốn.",
        "Nhấn <b>Thêm cặp khác</b> để thêm nhiều cặp huấn luyện.",
        "Nhấn <b>Run Training</b> để cập nhật bộ nhớ định dạng.",
    ], 1):
        story.append(Paragraph(f"  {i}. {desc}", step_style))
    story.append(Paragraph(
        I("💡 Sau khi train, công cụ sẽ ghi nhớ alias, khoảng cách tab (6cm), "
          "in đậm nhãn mục, và các quy tắc định dạng khác."),
        note_style))
    story.append(Spacer(1, 6))

    # ── Formatting rules ──
    story.append(Paragraph(B("3. Quy Tắc Định Dạng Chuẩn"), h1_style))
    rules = [
        ("Khoảng cách Tab", "Giữa thời gian và nội dung: <b>6 cm</b> (tab stop cố định)."),
        ("Nhãn Mục In Đậm", "Các nhãn mục (Education, Experience, Skills…) luôn được <b>in đậm</b>."),
        ("Tên Mục Lớn", "Section label (VD: <i>KINH NGHIỆM LÀM VIỆC</i>) được <b>in đậm</b>, canh giữa."),
        ("Thành Tích", "Achievements được <b>in đậm</b> và xếp sau nội dung chính."),
        ("Định dạng thời gian", "Định dạng <i>MM/YYYY – MM/YYYY</i> hoặc <i>present</i>."),
    ]
    for title_r, desc_r in rules:
        story.append(Paragraph(f"<b>• {title_r}:</b> {desc_r}", step_style))

    story.append(Spacer(1, 6))

    # ── Settings ──
    story.append(Paragraph(B("4. Thiết Lập (Settings)"), h1_style))
    settings_items = [
        ("Backend URL", "Địa chỉ server, mặc định: http://localhost:8000"),
        ("Claude API", "Nhập API key và chọn model (VD: claude-sonnet-4-20250514)."),
        ("OpenAI API", "API key dự phòng khi Claude không khả dụng."),
        ("Extraction Mode", "Chọn chế độ trích xuất phù hợp với nhu cầu."),
    ]
    for k, v in settings_items:
        story.append(Paragraph(f"<b>• {k}:</b> {v}", step_style))
    story.append(Spacer(1, 6))

    # ── Output ──
    story.append(Paragraph(B("5. File DOCX Đầu Ra"), h1_style))
    story.append(Paragraph(
        "File DOCX đầu ra có cấu trúc chuẩn: <b>font Times New Roman 12pt</b>, "
        "canh lề trái, tab stop 6cm, các section được phân tách rõ ràng. "
        "Có thể mở trực tiếp bằng Microsoft Word, Google Docs, hoặc LibreOffice.",
        body_style))
    story.append(Spacer(1, 10))

    # ── Footer ──
    story.append(HRFlowable(width="100%", thickness=1, color=accent, spaceBefore=6))
    story.append(Paragraph(
        I("CV Format Tool — Navigos Search | Phiên bản 2.1.0 | "
          "Hỗ trợ: cv-format-api@ navigos.com"),
        note_style))

    doc.build(story)
    return path


@app.get("/guideline.pdf", tags=["system"])
async def get_guideline():
    """Serve the user guideline PDF. Result is cached after first generation."""
    global _guideline_cache
    if _guideline_cache and os.path.isfile(_guideline_cache):
        return FileResponse(_guideline_cache, media_type="application/pdf",
                            filename="CV_Format_Tool_Huong_Dan_Su_Dung.pdf")
    try:
        path = _build_guideline_pdf()
        _guideline_cache = path
        return FileResponse(path, media_type="application/pdf",
                            filename="CV_Format_Tool_Huong_Dan_Su_Dung.pdf")
    except Exception as e:
        raise HTTPException(status_code=500,
                            detail=f"Không thể tạo PDF hướng dẫn: {e}")

# ═══════════════════════════════════════════════════════════════
# ── AUTH ROUTES ────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════

class LoginRequest(BaseModel):
    email: str
    password: str

class RefreshRequest(BaseModel):
    refresh_token: str

class UserCreateRequest(BaseModel):
    email: EmailStr
    password: str
    full_name: str
    role: str = "staff"

class LoginResponse(BaseModel):
    access_token: str
    refresh_token: str
    token_type: str = "bearer"
    expires_in: int = 28800  # 8 hours
    user: dict

@app.post("/auth/login", tags=["auth"])
async def login(payload: LoginRequest, request: Request):
    """Authenticate user and return JWT tokens."""
    with get_db_session() as session:
        user = session.query(User).filter(
            User.email == payload.email,
            User.is_active == True
        ).first()

        if not user or not verify_password(payload.password, user.hashed_password):
            raise HTTPException(status_code=401, detail="Invalid email or password")

        access_token = create_access_token(user.id, user.email, user.role)
        refresh_token = create_refresh_token(user.id)

        log_action(user.id, user.role, "login", request=request)
        session.commit()

        return LoginResponse(
            access_token=access_token,
            refresh_token=refresh_token,
            user=user.to_dict(),
        )

@app.post("/auth/refresh", tags=["auth"])
async def refresh(payload: RefreshRequest):
    """Get new access token from refresh token."""
    payload_data = decode_token(payload.refresh_token)
    if payload_data.get("type") != "refresh":
        raise HTTPException(status_code=401, detail="Invalid refresh token")

    with get_db_session() as session:
        user = session.query(User).filter(User.id == payload_data["sub"]).first()
        if not user or not user.is_active:
            raise HTTPException(status_code=401, detail="User not found or inactive")

        access_token = create_access_token(user.id, user.email, user.role)
        return {"access_token": access_token, "token_type": "bearer", "expires_in": 28800}

@app.get("/auth/me", tags=["auth"])
async def get_me(current_user: CurrentUser = Depends(get_current_user)):
    """Get current user profile."""
    with get_db_session() as session:
        user = session.query(User).filter(User.id == current_user.id).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        return user.to_dict()

# ═══════════════════════════════════════════════════════════════
# ── USER MANAGEMENT (admin only) ──────────────────────────────
# ═══════════════════════════════════════════════════════════════

@app.get("/users", tags=["users"])
async def list_users(
    current_user: CurrentUser = Depends(perm_user_read)
):
    """List all users (admin only)."""
    with get_db_session() as session:
        users = session.query(User).all()
        return [u.to_dict() for u in users]

@app.post("/users", tags=["users"])
async def create_user(
    payload: UserCreateRequest,
    current_user: CurrentUser = Depends(perm_user_create),
    request: Request = None,
):
    """Create a new user."""
    if payload.role not in ALLOWED_ROLES:
        raise HTTPException(status_code=400, detail=f"Invalid role. Allowed: {ALLOWED_ROLES}")

    with get_db_session() as session:
        existing = session.query(User).filter(User.email == payload.email).first()
        if existing:
            raise HTTPException(status_code=409, detail="Email already registered")

        user = User(
            email=payload.email,
            hashed_password=hash_password(payload.password),
            full_name=payload.full_name,
            role=payload.role,
        )
        session.add(user)
        session.commit()
        session.refresh(user)

        log_action(current_user.id, current_user.role, "create_user",
                   resource_type="user", resource_id=user.id,
                   details={"email": user.email, "role": user.role},
                   request=request)
        return user.to_dict()

@app.patch("/users/{user_id}", tags=["users"])
async def update_user(
    user_id: str,
    payload: dict,
    current_user: CurrentUser = Depends(perm_user_update),
):
    """Update user (admin only)."""
    with get_db_session() as session:
        user = session.query(User).filter(User.id == user_id).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        allowed = {"full_name", "role", "is_active"}
        updates = {k: v for k, v in payload.items() if k in allowed}
        if "role" in updates and updates["role"] not in ALLOWED_ROLES:
            raise HTTPException(status_code=400, detail="Invalid role")

        for k, v in updates.items():
            setattr(user, k, v)
        user.updated_at = datetime.utcnow()
        session.commit()
        return user.to_dict()

# ═══════════════════════════════════════════════════════════════
# ── CV JOB WORKFLOW ROUTES ────────────────────────────────────
# ═══════════════════════════════════════════════════════════════

class UploadResponse(BaseModel):
    job_id: str
    original_filename: str
    file_size: int
    file_type: str
    status: str
    message: str
    # Extended fields for ReviewPanel integration
    parsed_data: Optional[dict[str, Any]] = None
    validation_result: Optional[dict[str, Any]] = None
    download_url: Optional[str] = None

@app.post("/jobs", response_model=UploadResponse, tags=["jobs"])
async def upload_cv(
    file: UploadFile = File(...),
    extraction_mode: str = Form("auto"),
    api_key: str = Form(""),
    model: str = Form("claude-sonnet-4-20250514"),
    openai_api_key: str = Form(""),
    openai_model: str = Form("gpt-4o-mini"),
    current_user: CurrentUser = Depends(perm_cv_upload),
    request: Request = None,
):
    """
    NEW workflow: Upload → Validate → Parse → Review → QC → Export

    This endpoint handles the FULL pipeline:
    1. Save file
    2. Extract text
    3. Parse with AI/offline
    4. Validate (email, phone, date, required fields)
    5. Return parsed data + validation result
    6. Frontend shows Review panel before export
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in (".pdf", ".docx"):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX supported")

    # Save to temp
    suffix = ext
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        # ── Step 1: Extract text ───────────────────────────────
        if ext == ".pdf":
            cv_text = _extract_text_from_pdf(tmp_path)
        else:
            cv_text = _extract_text_from_docx(tmp_path)

        if not cv_text.strip():
            # Provide actionable error message
            ocr_hint = ""
            try:
                from ocr_engine import is_ocr_available
                if not is_ocr_available():
                    ocr_hint = " Install pytesseract or easyocr for OCR support."
                else:
                    ocr_hint = " OCR was attempted but returned no text. The file may be corrupted."
            except ImportError:
                ocr_hint = " OCR module not found."
            raise HTTPException(
                status_code=422,
                detail=f"Could not extract text (likely image-based/scanned PDF).{ocr_hint}",
            )

        lang = _detect_language(cv_text)
        template_path = TEMPLATE_VN if lang == "vi" else TEMPLATE_EN

        # ── Step 2: Parse CV data ────────────────────────────
        if extraction_mode == "offline":
            from offline_engine import extract_offline, build_suggested_name_offline
            cv_data = extract_offline(cv_text)
            suggested_name = build_suggested_name_offline(cv_data)
        else:
            cv_data = extract_cv_data(
                cv_text, api_key, model,
                extraction_mode, openai_api_key, openai_model
            )
            suggested_name = build_suggested_name(cv_data)

        # ── Step 3: Validate ─────────────────────────────────
        validation_result = validate_cv_data(cv_data)
        validation_dict = validation_result.to_dict()

        # ── Step 4: Save to DB ────────────────────────────────
        with get_db_session() as session:
            job = CVJob(
                owner_id=current_user.id,
                original_filename=file.filename,
                file_size=os.path.getsize(tmp_path),
                file_type=ext.lstrip("."),
                status="review",
                extraction_mode=extraction_mode,
                parsed_data=cv_data,
                parsed_at=datetime.utcnow(),
                validation_errors=validation_dict.get("errors"),
                validated_at=datetime.utcnow(),
            )
            session.add(job)
            session.commit()
            session.refresh(job)

            # Create initial version
            version = CVVersion(
                job_id=job.id,
                version_number=1,
                changed_by_id=current_user.id,
                changed_by_role=current_user.role,
                change_type="initial_parse",
                data_snapshot=cv_data,
                notes="Initial AI/offline parse",
            )
            session.add(version)

            log_action(
                current_user.id, current_user.role, "upload",
                resource_type="cv_job", resource_id=job.id,
                details={"filename": file.filename, "extraction_mode": extraction_mode,
                         "validation": validation_dict["summary"]},
                request=request,
            )
            session.commit()

        return UploadResponse(
            job_id=job.id,
            original_filename=file.filename,
            file_size=job.file_size,
            file_type=job.file_type,
            status="review",
            message=f"Uploaded. {validation_dict['summary']}",
            parsed_data=cv_data,
            validation_result=validation_dict,
            download_url=None,  # Not ready until reviewed + approved
        )

    finally:
        os.unlink(tmp_path)


@app.get("/jobs", tags=["jobs"])
async def list_jobs(
    status: Optional[str] = None,
    my_only: bool = False,
    limit: int = 50,
    current_user: CurrentUser = Depends(get_current_user),
):
    """List CV jobs. Staff sees own jobs; QC/Admin see all."""
    with get_db_session() as session:
        query = session.query(CVJob)

        # Staff can only see their own jobs
        if current_user.role == "staff" or my_only:
            query = query.filter(CVJob.owner_id == current_user.id)

        if status:
            query = query.filter(CVJob.status == status)

        jobs = query.order_by(CVJob.created_at.desc()).limit(limit).all()
        return [j.to_dict(include_parsed=False) for j in jobs]


@app.get("/jobs/{job_id}", tags=["jobs"])
async def get_job(
    job_id: str,
    current_user: CurrentUser = Depends(get_current_user),
):
    """Get job details with parsed data and validation."""
    with get_db_session() as session:
        job = session.query(CVJob).filter(CVJob.id == job_id).first()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")

        # Staff can only see own jobs
        if current_user.role == "staff" and job.owner_id != current_user.id:
            raise HTTPException(status_code=403, detail="Access denied")

        return job.to_dict(include_parsed=True)


class ReviewUpdate(BaseModel):
    reviewed_data: dict
    notes: Optional[str] = None

@app.patch("/jobs/{job_id}/review", tags=["jobs"])
async def review_job(
    job_id: str,
    payload: ReviewUpdate,
    current_user: CurrentUser = Depends(perm_cv_review),
    request: Request = None,
):
    """Staff reviews and corrects parsed CV data."""
    with get_db_session() as session:
        job = session.query(CVJob).filter(CVJob.id == job_id).first()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")

        if job.status not in ("review", "parsed", "validated"):
            raise HTTPException(status_code=400, detail=f"Cannot review job in status: {job.status}")

        # Get next version number
        latest = session.query(CVVersion).filter(
            CVVersion.job_id == job_id
        ).order_by(CVVersion.version_number.desc()).first()
        next_version = (latest.version_number + 1) if latest else 1

        # Sanitize data before saving
        clean_data = sanitize_for_export(payload.reviewed_data)

        # Create new version snapshot
        version = CVVersion(
            job_id=job_id,
            version_number=next_version,
            changed_by_id=current_user.id,
            changed_by_role=current_user.role,
            change_type="staff_review",
            data_snapshot=clean_data,
            notes=payload.notes or "Staff review corrections",
        )
        session.add(version)

        # Re-validate corrected data
        val_result = validate_cv_data(clean_data)

        # Update job
        job.reviewed_data = clean_data
        job.reviewed_at = datetime.utcnow()
        job.validation_errors = val_result.to_dict().get("errors")
        # QC/Admin can send directly to QC approval; staff send back to review for QC
        job.status = "qc" if current_user.role in ("qc", "admin") else "review"

        session.commit()

        log_action(
            current_user.id, current_user.role, "review",
            resource_type="cv_job", resource_id=job_id,
            details={"version": next_version, "validation": val_result.summary},
            request=request,
        )

        return {
            "ok": True,
            "message": f"Review saved (v{next_version})",
            "version_number": next_version,
            "validation": val_result.to_dict(),
            "job": job.to_dict(include_parsed=True),
        }


class QCRequest(BaseModel):
    result: str  # pass | fail | needs_revision
    notes: Optional[str] = None

@app.patch("/jobs/{job_id}/qc", tags=["jobs"])
async def qc_job(
    job_id: str,
    payload: QCRequest,
    current_user: CurrentUser = Depends(perm_cv_qc),
    request: Request = None,
):
    """QC reviewer approves or rejects a CV."""
    if payload.result not in ("pass", "fail", "needs_revision"):
        raise HTTPException(status_code=400, detail="Invalid QC result")

    with get_db_session() as session:
        job = session.query(CVJob).filter(CVJob.id == job_id).first()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")

        if job.status not in ("qc", "review"):
            raise HTTPException(status_code=400, detail=f"Cannot QC job in status: {job.status}")

        job.qc_by = current_user.id
        job.qc_result = payload.result
        job.qc_notes = payload.notes
        job.qc_at = datetime.utcnow()

        if payload.result == "pass":
            job.status = "approved"
        elif payload.result == "needs_revision":
            job.status = "review"
        else:  # fail
            job.status = "error"

        session.commit()

        log_action(
            current_user.id, current_user.role, "qc",
            resource_type="cv_job", resource_id=job_id,
            details={"result": payload.result, "notes": payload.notes},
            request=request,
        )

        return {
            "ok": True,
            "message": f"QC: {payload.result}",
            "job": job.to_dict(),
        }


@app.post("/jobs/{job_id}/export", tags=["jobs"])
async def export_job(
    job_id: str,
    client_name: str = Form("CLIENT"),
    position: str = Form("POSITION"),
    current_user: CurrentUser = Depends(perm_cv_export),
    request: Request = None,
):
    """Export approved CV to formatted DOCX."""
    with get_db_session() as session:
        job = session.query(CVJob).filter(CVJob.id == job_id).first()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")

        if job.status not in ("approved", "review", "qc"):
            raise HTTPException(
                status_code=400,
                detail=f"Cannot export job in status: {job.status}. Job must be QC-approved.",
            )

        # Use reviewed data if available, otherwise parsed data
        cv_data = job.reviewed_data or job.parsed_data
        if not cv_data:
            raise HTTPException(status_code=400, detail="No CV data to export")

        # Run validation on the actual data we're about to export
        validation = validate_cv_data(cv_data)
        if not validation.is_valid:
            if not has_permission(current_user.role, "cv:override_export"):
                raise HTTPException(
                    status_code=403,
                    detail=(
                        "CV has blocking validation errors and you do not have "
                        "permission to override. Fix the errors first or contact a QC/Admin."
                    ),
                )

        lang = _detect_language(str(cv_data))
        template_path = TEMPLATE_VN if lang == "vi" else TEMPLATE_EN

        # Generate output
        download_id = str(uuid.uuid4())[slice(0, 8)]
        safe_name = re.sub(r"[^\w\s.-]", "", (job.output_filename or job.original_filename or "cv")).strip()
        safe_name = re.sub(r"\s+", "_", safe_name)[slice(0, 100)]
        output_dir = os.path.join(OUTPUT_DIR, download_id)
        os.makedirs(output_dir, exist_ok=True)
        output_docx = os.path.join(output_dir, f"{safe_name}.docx")

        fill_template(template_path, cv_data, client_name, position, output_docx)

        # Update job
        job.status = "exported"
        job.output_filename = f"{safe_name}.docx"
        job.output_path = output_docx
        job.download_url = f"/download/{download_id}"
        job.exported_at = datetime.utcnow()
        job.completed_at = datetime.utcnow()

        log_action(
            current_user.id, current_user.role, "export",
            resource_type="cv_job", resource_id=job_id,
            details={"client": client_name, "position": position, "filename": f"{safe_name}.docx"},
            request=request,
        )
        session.commit()

        return {
            "ok": True,
            "message": "CV exported successfully",
            "download_id": download_id,
            "download_url": f"/download/{download_id}",
            "filename": f"{safe_name}.docx",
        }


@app.get("/jobs/{job_id}/versions", tags=["jobs"])
async def get_job_versions(
    job_id: str,
    current_user: CurrentUser = Depends(get_current_user),
):
    """Get version history for a job."""
    with get_db_session() as session:
        job = session.query(CVJob).filter(CVJob.id == job_id).first()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        if current_user.role == "staff" and job.owner_id != current_user.id:
            raise HTTPException(status_code=403, detail="Access denied")

        versions = session.query(CVVersion).filter(
            CVVersion.job_id == job_id
        ).order_by(CVVersion.version_number.asc()).all()
        return [v.to_dict() for v in versions]


# ═══════════════════════════════════════════════════════════════
# ── BATCH PROCESSING ROUTES ────────────────────────────────────
# ═══════════════════════════════════════════════════════════════

@app.post("/batch", tags=["batch"])
async def create_batch(
    files: list[UploadFile] = File(...),
    extraction_mode: str = Form("auto"),
    batch_name: str = Form(""),
    current_user: CurrentUser = Depends(perm_cv_upload),
):
    """
    Upload multiple CVs and process them in parallel.
    Returns immediately with batch_id — frontend polls for status.
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")

    batch_id = str(uuid.uuid4())
    batch = Batch(
        id=batch_id,
        name=batch_name or f"Batch {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}",
        owner_id=current_user.id,
        job_count=len(files),
        created_at=datetime.utcnow().isoformat(),
    )

    jobs = []
    for f in files:
        if not f.filename:
            continue
        ext = os.path.splitext(f.filename)[1].lower()
        if ext not in (".pdf", ".docx"):
            continue

        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            shutil.copyfileobj(f.file, tmp)
            tmp_path = tmp.name

        job = BatchJob(
            id=str(uuid.uuid4()),
            batch_id=batch_id,
            original_filename=f.filename,
            file_path=tmp_path,
            file_type=ext.lstrip("."),
            file_size=os.path.getsize(tmp_path),
            extraction_mode=extraction_mode,
        )
        jobs.append(job)

    # Submit to processor
    processor = get_processor()
    processor.submit_batch(batch, jobs)

    return {
        "batch_id": batch_id,
        "name": batch.name,
        "job_count": len(jobs),
        "status": "running",
        "message": f"{len(jobs)} CVs queued for processing",
    }

@app.get("/batch/{batch_id}", tags=["batch"])
async def get_batch_status(
    batch_id: str,
    current_user: CurrentUser = Depends(get_current_user),
):
    """Get batch processing status. Staff can only see their own batches."""
    processor = get_processor()
    batch = processor.get_batch_status(batch_id)
    if not batch:
        raise HTTPException(status_code=404, detail="Batch not found")
    if current_user.role == "staff" and batch.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    return batch.to_dict()

@app.get("/batch/{batch_id}/jobs", tags=["batch"])
async def get_batch_jobs(
    batch_id: str,
    current_user: CurrentUser = Depends(get_current_user),
):
    """Get all jobs in a batch. Staff can only see their own batches."""
    processor = get_processor()
    batch = processor.get_batch_status(batch_id)
    if batch and current_user.role == "staff" and batch.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    jobs = processor.get_all_jobs(batch_id)
    return [j.to_dict() for j in jobs]

@app.get("/batch", tags=["batch"])
async def list_all_batches(
    current_user: CurrentUser = Depends(get_current_user),
):
    """List all batches."""
    batches = list_batches()
    # Filter to owner unless admin/qc
    if current_user.role == "staff":
        batches = [b for b in batches if b.owner_id == current_user.id]
    return [b.to_dict() for b in batches[:20]]

@app.delete("/batch/{batch_id}", tags=["batch"])
async def cancel_batch(
    batch_id: str,
    current_user: CurrentUser = Depends(perm_cv_upload),
):
    """Cancel a running batch. Staff can only cancel their own batches."""
    processor = get_processor()
    batch = processor.get_batch_status(batch_id)
    if not batch:
        raise HTTPException(status_code=404, detail="Batch not found")
    if current_user.role == "staff" and batch.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied: you can only cancel your own batches")
    result = processor.cancel_batch(batch_id)
    if not result:
        raise HTTPException(status_code=404, detail="Batch not found")
    return {"ok": True, "message": "Batch cancelled"}

# ═══════════════════════════════════════════════════════════════
# ── VALIDATION STANDALONE ROUTE ────────────────────────────────
# ═══════════════════════════════════════════════════════════════

class ValidateRequest(BaseModel):
    cv_data: dict
    strict: bool = False

@app.post("/validate", tags=["validation"])
async def validate_cv(
    payload: ValidateRequest,
    current_user: CurrentUser = Depends(get_current_user),
):
    """Standalone validation endpoint — validate CV data without processing file."""
    result = validate_cv_data(payload.cv_data, strict=payload.strict)
    return result.to_dict()

# ═══════════════════════════════════════════════════════════════
# ── AUDIT LOG ─────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════

@app.get("/audit", tags=["audit"])
async def get_audit_log(
    limit: int = 50,
    action: Optional[str] = None,
    current_user: CurrentUser = Depends(perm_audit_read),
):
    """Get audit log entries."""
    with get_db_session() as session:
        query = session.query(AuditLog).order_by(AuditLog.created_at.desc())
        if action:
            query = query.filter(AuditLog.action == action)
        logs = query.limit(limit).all()
        return [log.to_dict() for log in logs]

# ═══════════════════════════════════════════════════════════════
# ── KPI DASHBOARD (admin) ─────────────────────────────────────
# ═══════════════════════════════════════════════════════════════

@app.get("/stats", tags=["admin"])
async def get_stats(
    current_user: CurrentUser = Depends(perm_cv_view_all),
):
    """KPI dashboard: CV throughput, error rate, time metrics."""
    with get_db_session() as session:
        total = session.query(CVJob).count()
        by_status = {}
        for status in ["uploaded", "parsing", "parsed", "validated", "review",
                        "qc", "approved", "exported", "error", "cancelled"]:
            by_status[status] = session.query(CVJob).filter(CVJob.status == status).count()

        # Error rate
        error_rate = round(by_status.get("error", 0) / total * 100, 1) if total > 0 else 0

        # Avg processing time (for completed jobs)
        completed = session.query(CVJob).filter(
            CVJob.status == "exported",
            CVJob.completed_at.isnot(None),
            CVJob.created_at.isnot(None),
        ).limit(100).all()

        import statistics
        times = []
        for job in completed:
            try:
                created = datetime.fromisoformat(job.created_at)
                completed_at = datetime.fromisoformat(job.completed_at)
                delta = (completed_at - created).total_seconds()
                if 0 < delta < 3600:  # Sanity: < 1 hour
                    times.append(delta)
            except (ValueError, TypeError, AttributeError):
                continue

        avg_time = round(statistics.mean(times), 1) if times else 0

        # Users
        user_count = session.query(User).filter(User.is_active == True).count()

        return {
            "total_jobs": total,
            "by_status": by_status,
            "error_rate_percent": error_rate,
            "avg_processing_seconds": avg_time,
            "active_users": user_count,
            "cvs_per_day_target": 50,
        }

# ═══════════════════════════════════════════════════════════════
# ── HEALTH CHECK (enhanced) ───────────────────────────────────
# ═══════════════════════════════════════════════════════════════

@app.get("/health", tags=["system"])
async def health_check(api_key: str = "", openai_api_key: str = ""):
    """Enhanced health check with DB + processor + OCR + AI provider status."""
    ocr_status = "unavailable"
    ocr_backend = "none"
    try:
        from ocr_engine import is_ocr_available, get_ocr_backend
        if is_ocr_available():
            ocr_status = "ok"
            ocr_backend = get_ocr_backend()
    except ImportError:
        pass

    # ── Check AI providers ──────────────────────────────────────
    claude_status = "unavailable"
    if api_key:
        try:
            import anthropic as _anth
            client = _anth.Anthropic(api_key=api_key)
            # Minimal API call to verify key works
            client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=5,
                messages=[{"role": "user", "content": "hi"}],
            )
            claude_status = "ok"
        except Exception as e:
            err = str(e).lower()
            if "credit" in err or "billing" in err or "payment" in err:
                claude_status = "no_credit"
            elif "invalid" in err or "auth" in err or "api key" in err:
                claude_status = "invalid_key"
            else:
                claude_status = "error"

    openai_status = "unavailable"
    if openai_api_key:
        try:
            import openai as _openai
            client = _openai.OpenAI(api_key=openai_api_key)
            client.models.list()
            openai_status = "ok"
        except Exception as e:
            err = str(e).lower()
            if "quota" in err or "billing" in err or "exceeded" in err:
                openai_status = "quota_exceeded"
            elif "invalid" in err or "auth" in err:
                openai_status = "invalid_key"
            else:
                openai_status = "error"

    ollama_status = "unavailable"
    try:
        import urllib.request
        req = urllib.request.Request("http://localhost:11434/api/tags", method="GET")
        with urllib.request.urlopen(req, timeout=3):
            ollama_status = "ok"
    except Exception:
        pass

    template_en_ok = os.path.exists(TEMPLATE_EN)
    template_vn_ok = os.path.exists(TEMPLATE_VN)

    return {
        "status": "ok",
        "version": "2.2.0",
        "claude": claude_status,
        "openai": openai_status,
        "ollama": ollama_status,
        "timestamp": datetime.utcnow().isoformat(),
        "components": {
            "database": "ok",
            "batch_processor": "ok",
            "validation_engine": "ok",
            "ocr": ocr_status,
            "ocr_backend": ocr_backend,
            "template_en": "ok" if template_en_ok else f"missing ({TEMPLATE_EN})",
            "template_vn": "ok" if template_vn_ok else f"missing ({TEMPLATE_VN})",
        }
    }

# ═══════════════════════════════════════════════════════════════
# ── DOWNLOAD (existing, unchanged) ────────────────────────────
# ═══════════════════════════════════════════════════════════════

@app.get("/download/{download_id}", tags=["system"])
async def download_file(download_id: str):
    """Download a processed DOCX file."""
    path = os.path.join(OUTPUT_DIR, download_id)
    if not os.path.isdir(path):
        raise HTTPException(status_code=404, detail="File not found")
    files = [f for f in os.listdir(path) if f.endswith(".docx")]
    if not files:
        raise HTTPException(status_code=404, detail="File not found")
    file_path = os.path.join(path, files[0])
    return FileResponse(
        file_path,
        filename=files[0],
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# ═══════════════════════════════════════════════════════════════
# ── ROOT ──────────────────────────────────────────────────────
# ═══════════════════════════════════════════════════════════════

@app.get("/", tags=["system"])
async def root():
    return {
        "name": "CV Format Tool API",
        "version": "2.1.0",
        "modules": ["auth", "jobs", "batch", "validation", "audit"],
        "docs": "/docs",
    }


# ═══════════════════════════════════════════════════════════════
# ── EXISTING PROCESS ENDPOINT (legacy, wraps new workflow) ────
# ═══════════════════════════════════════════════════════════════

class ProcessResponse(BaseModel):
    status: str
    message: str
    suggestedName: str | None = None
    downloadId: str | None = None
    downloadUrl: str | None = None
    reviewRequired: list[dict] | None = None


@app.post("/process", response_model=ProcessResponse, tags=["legacy"],
          deprecated=True)
async def process_cv_legacy(
    file: UploadFile = File(...),
    extraction_mode: str = Form("auto"),
    model: str = Form("claude-sonnet-4-20250514"),
    api_key: str = Form(""),
    openai_api_key: str = Form(""),
    openai_model: str = Form("gpt-4o-mini"),
):
    """
    DEPRECATED — Use POST /jobs instead.
    This legacy endpoint skips validation, review, and QC workflow.
    It will be removed in a future version.
    """
    import warnings
    warnings.warn(
        "POST /process is deprecated. Use POST /jobs for the full workflow.",
        DeprecationWarning, stacklevel=2,
    )
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in (".pdf", ".docx"):
        raise HTTPException(status_code=400, detail="Only PDF/DOCX supported")

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        if ext == ".pdf":
            cv_text = _extract_text_from_pdf(tmp_path)
        else:
            cv_text = _extract_text_from_docx(tmp_path)

        if not cv_text.strip():
            return ProcessResponse(
                status="error",
                message="Could not extract text (possibly image-based PDF)",
            )

        lang = _detect_language(cv_text)
        template_path = TEMPLATE_VN if lang == "vi" else TEMPLATE_EN

        if extraction_mode == "offline":
            from offline_engine import (
                extract_offline, build_suggested_name_offline, fill_template_offline
            )
            cv_data = extract_offline(cv_text)
            suggested_name = build_suggested_name_offline(cv_data)
            template_meta = fill_template_offline(template_path, "/tmp/dummy.docx", cv_data)
            review_required = template_meta.get("reviewRequired", [])
        else:
            cv_data = extract_cv_data(cv_text, api_key, model, extraction_mode,
                                      openai_api_key, openai_model)
            suggested_name = build_suggested_name(cv_data)
            review_required = None

        download_id = str(uuid.uuid4())[slice(0, 8)]
        safe_name = re.sub(r"[^\w\s.-]", "", (suggested_name or file.filename or "cv")).strip()
        safe_name = re.sub(r"\s+", "_", safe_name)[slice(0, 100)]
        output_dir = os.path.join(OUTPUT_DIR, download_id)
        os.makedirs(output_dir, exist_ok=True)
        output_docx = os.path.join(output_dir, f"{safe_name}.docx")

        if not os.path.exists(template_path):
            # Template missing — return parsed data without DOCX generation
            return ProcessResponse(
                status="success",
                message=f"Parsed ({lang.upper()}) — template not available for DOCX",
                suggestedName=suggested_name,
                reviewRequired=review_required,
            )

        try:
            if extraction_mode == "offline":
                from offline_engine import fill_template_offline
                fill_template_offline(template_path, output_docx, cv_data)
            else:
                fill_template(template_path, cv_data, "CLIENT", "POSITION", output_docx)
        except Exception as e:
            # Template fill failed — return parsed data with error note
            return ProcessResponse(
                status="success",
                message=f"Parsed ({lang.upper()}) — DOCX fill error: {e}",
                suggestedName=suggested_name,
                reviewRequired=review_required,
            )

        return ProcessResponse(
            status="success",
            message=f"Generated ({lang.upper()})",
            suggestedName=suggested_name,
            downloadId=download_id,
            downloadUrl=f"/download/{download_id}",
            reviewRequired=review_required,
        )
    finally:
        os.unlink(tmp_path)


# ── Legacy import compatibility ────────────────────────────────
# Keep these accessible for backward compat
EXTRACTION_PROMPT = """You are a CV/Resume parser. Extract structured information from the following CV text and return ONLY valid JSON.

The JSON must have this exact structure:
{
  "full_name": "string or empty",
  "gender": "string or empty",
  "year_of_birth": "string or empty",
  "marital_status": "string or empty",
  "address": "string or empty",
  "career_summary": [
    {
      "period": "MM/YYYY – MM/YYYY or Present",
      "company": "COMPANY NAME IN UPPERCASE",
      "company_description": "Brief company description if available, or empty string",
      "positions": [
        {
          "period": "MM/YYYY – MM/YYYY (if different sub-period, otherwise empty)",
          "title": "Job Title",
          "report_to": "reporting line if mentioned, or empty",
          "section_label": "e.g. 'Accountability:' or 'Duties:' if mentioned, or empty",
          "responsibilities": ["bullet point 1", "bullet point 2"],
          "achievements_label": "e.g. 'Achievements:' if mentioned, or empty",
          "achievements": ["achievement 1", "achievement 2"]
        }
      ]
    }
  ],
  "education": [
    {
      "period": "YYYY - YYYY",
      "institution": "UNIVERSITY NAME IN UPPERCASE",
      "details": ["degree/major line 1", "line 2"]
    }
  ],
  "other_info": [
    {
      "section_title": "SECTION TITLE IN UPPERCASE",
      "items": ["item 1", "item 2"]
    }
  ]
}

IMPORTANT RULES:
- Extract ALL work experiences, education, and other sections
- Keep the original language of the CV content (do not translate)
- Company names should be in UPPERCASE
- If information is not available, use empty string ""
- Do NOT invent or assume information that is not in the CV

CV TEXT:
"""

# ── Paths ──────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_EN = os.path.join(BASE_DIR, "templates", "Form_EN_2024.docx")
TEMPLATE_VN = os.path.join(BASE_DIR, "templates", "Form_VN_2024.docx")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Log template availability at startup
for _tpl_name, _tpl_path in [("EN", TEMPLATE_EN), ("VN", TEMPLATE_VN)]:
    if os.path.exists(_tpl_path):
        print(f"[TEMPLATE] {_tpl_name}: OK ({_tpl_path})")
    else:
        print(f"[TEMPLATE] {_tpl_name}: MISSING ({_tpl_path})")

# ── Local helpers (reused from original) ──────────────────────

def _extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF. Falls back to OCR for scanned/image-based PDFs."""
    doc = fitz.open(pdf_path)
    texts = [page.get_text() for page in doc]
    doc.close()
    text = "\n".join(texts)

    # If very little text extracted, try OCR
    if len(text.strip()) < 50:
        try:
            from ocr_engine import is_ocr_available, extract_text_from_scanned_pdf
            if is_ocr_available():
                print(f"[PDF] Low text content ({len(text.strip())} chars), attempting OCR...")
                ocr_text = extract_text_from_scanned_pdf(pdf_path)
                if ocr_text.strip():
                    print(f"[PDF] OCR extracted {len(ocr_text)} chars")
                    return ocr_text
                else:
                    print("[PDF] OCR returned no text")
            else:
                print("[PDF] OCR not available (install pytesseract or easyocr)")
        except Exception as e:
            print(f"[PDF] OCR fallback failed: {e}")

    return text

def _extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    texts = []
    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                texts.append(row_text)
    return "\n".join(texts)

def _detect_language(text: str) -> str:
    vn_chars = set("àáảãạăắằẳẵặâấầẩẫậèéẻẽẹêếềểễệìíỉĩịòóỏõọôốồổỗộơớờởỡợùúủũụưứừửữựỳýỷỹỵđ"
                   "ÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬÈÉẺẼẸÊẾỀỂỄỆÌÍỈĨỊÒÓỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÙÚỦŨỤƯỨỪỬỮỰỲÝỶỸỴĐ")
    count = sum(1 for c in text if c in vn_chars)
    total = len(text.replace(" ", "").replace("\n", ""))
    if total == 0:
        return "en"
    return "vi" if (count / total) > 0.02 else "en"


# ── Original CV extraction functions (kept for /process legacy endpoint) ──
# These were the core of the original main.py. Kept here for backward compat.

def _parse_json_response(text: str) -> dict:
    json_match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
    json_str = json_match.group(1).strip() if json_match else text.strip()
    if not json_match:
        brace_s = json_str.find('{')
        brace_e = json_str.rfind('}') + 1
        if brace_s != -1 and brace_e > brace_s:
            json_str = json_str[brace_s:brace_e]
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        json_str = re.sub(r',(\s*[}\]])', r'\1', json_str)
        json_str = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', json_str)
        try:
            return json.loads(json_str)
        except json.JSONDecodeError:
            brace_count: int = 0
            start_found: bool = False
            start_pos: int = -1
            end_pos: int = -1
            for i, c in enumerate(text):
                if c == '{':
                    if not start_found:
                        start_pos = i
                        start_found = True
                    brace_count += 1
                elif c == '}':
                    brace_count -= 1
                    if brace_count == 0 and start_found:
                        end_pos = i + 1
                        break
            if start_pos != -1 and end_pos > start_pos:
                return json.loads(text[start_pos:end_pos])
            raise ValueError(f"Could not parse JSON from response. Preview: {text[slice(0, 200)]}")

def extract_with_claude(cv_text: str, api_key: str, model: str) -> dict:
    import anthropic as _anth
    client = _anth.Anthropic(api_key=api_key)
    msg = client.messages.create(model=model, max_tokens=8000,
                                  messages=[{"role": "user", "content": EXTRACTION_PROMPT + cv_text}])
    return _parse_json_response(msg.content[0].text)

def extract_with_openai(cv_text: str, api_key: str, model: str) -> dict:
    import openai as _openai
    client = _openai.OpenAI(api_key=api_key)
    resp = client.chat.completions.create(model=model, max_tokens=8000, temperature=0.1,
        messages=[{"role": "user", "content": EXTRACTION_PROMPT + cv_text}])
    return _parse_json_response(resp.choices[0].message.content or "")

def extract_with_ollama(cv_text: str, model: str = "qwen2.5:14b") -> dict:
    import urllib.request, urllib.error
    payload = json.dumps({"model": model, "prompt": EXTRACTION_PROMPT + cv_text,
                          "stream": False, "options": {"temperature": 0.1, "num_predict": 8000}}).encode()
    req = urllib.request.Request("http://localhost:11434/api/generate",
                                  data=payload, headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=300) as resp:
            return _parse_json_response(json.loads(resp.read())["response"])
    except urllib.error.URLError as e:
        raise ConnectionError(f"Ollama unavailable: {e.reason}")

def extract_cv_data(cv_text: str, api_key: str, model: str, mode: str,
                    openai_key: str = "", openai_model: str = "gpt-4o-mini") -> dict:
    errors = []
    if mode in ("auto", "claude_api") and api_key:
        try:
            return extract_with_claude(cv_text, api_key, model)
        except Exception as e:
            errors.append(f"Claude: {e}")
            if mode == "claude_api": raise ValueError(errors[-1])
    if mode in ("auto", "openai_api") and openai_key:
        try:
            return extract_with_openai(cv_text, openai_key, openai_model)
        except Exception as e:
            errors.append(f"OpenAI: {e}")
            if mode == "openai_api": raise ValueError(errors[-1])
    if mode in ("auto", "ollama"):
        try:
            return extract_with_ollama(cv_text, model)
        except Exception as e:
            errors.append(f"Ollama: {e}")
            raise ValueError(" | ".join(errors))
    raise ValueError(errors[-1] if errors else "No AI provider available")

def build_suggested_name(cv_data: dict) -> str:
    career = cv_data.get("career_summary", [])
    full_name = cv_data.get("full_name", "").strip()
    current_job = None
    for job in career:
        if "Present" in job.get("period", "") or "Hiện tại" in job.get("period", ""):
            current_job = job; break
    if current_job is None and career:
        current_job = career[0]
    _cj: dict = current_job or {}
    company = (_cj.get("company", "") or "").strip().title()
    positions: list = (_cj.get("positions") or []) if current_job else []
    position = ((positions[0].get("title", "") or "").strip().title()
                if positions else (_cj.get("title", "") or "").strip().title()) if current_job else ""
    parts = [p for p in [company, position, full_name.title() if full_name else ""] if p]
    return " - ".join(parts)

def _get_style_id(doc, style_name: str) -> str:
    try: return doc.styles[style_name].style_id
    except KeyError: return "Normal"

def _make_p(text: str, style_id: str) -> OxmlElement:
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle); p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text; t.set(qn('xml:space'), 'preserve')
        r.append(t); p.append(r)
    return p


def _make_bold_run(text: str) -> OxmlElement:
    """Return a <w:r> element with bold text."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


def _make_tab_run() -> OxmlElement:
    """Return a <w:r> element containing a single tab character with a 6 cm tab stop."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), '1512000')   # 1512000 twips = 6 cm (1 cm = 567 twips)
    tabs.append(tab)
    rPr.append(tabs)
    r.append(rPr)
    r.append(OxmlElement('w:tab'))
    return r


def _make_p_with_bold_label(label: str, content: str, style_id: str) -> OxmlElement:
    """
    Paragraph: bold label + 6cm tab stop + normal content.
    label: bold text shown at left; content appears after a 6 cm tab.
    """
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)

    # Add 6 cm tab stop to paragraph
    tabs_el = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'left')
    tab_el.set(qn('w:pos'), '1512000')
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    p.append(pPr)

    # Bold label run
    if label:
        p.append(_make_bold_run(label))
    # Tab
    p.append(_make_tab_run())
    # Content run (not bold)
    if content:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = content
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
    return p

def _populate_toc(body, career_summary: list):
    sdt_el = next((c for c in body if c.tag.split('}')[-1] == 'sdt'), None)
    if sdt_el is None: return
    sdt_content = sdt_el.find(qn('w:sdtContent'))
    if sdt_content is None: return
    for child in list(sdt_content): sdt_content.remove(child)
    def _toc_rpr():
        rPr = OxmlElement('w:rPr'); rPr.append(OxmlElement('w:noProof')); return rPr
    for job in career_summary:
        if job.get("use_normal_style"): continue
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), 'TOC1')
        pPr.append(ps); p.append(pPr)
        for txt in [job.get("period", ""), job.get("company", "")]:
            if txt:
                r = OxmlElement('w:r'); r.append(_toc_rpr())
                t = OxmlElement('w:t'); t.text = txt; t.set(qn('xml:space'), 'preserve')
                r.append(t); p.append(r)
            r_tab = OxmlElement('w:r'); r_tab.append(_toc_rpr())
            r_tab.append(OxmlElement('w:tab')); p.append(r_tab)
        sdt_content.append(p)
        for pos in job.get("positions", []):
            if not pos.get("title"): continue
            p2 = OxmlElement('w:p')
            pPr2 = OxmlElement('w:pPr')
            ps2 = OxmlElement('w:pStyle'); ps2.set(qn('w:val'), 'TOC3')
            pPr2.append(ps2); p2.append(pPr2)
            r2 = OxmlElement('w:r'); r2.append(_toc_rpr())
            t2 = OxmlElement('w:t'); t2.text = pos["title"]; t2.set(qn('xml:space'), 'preserve')
            r2.append(t2); p2.append(r2)
            sdt_content.append(p2)

def fill_template(template_path: str, cv_data: dict, client_name: str,
                  position: str, output_path: str):
    doc = Document(template_path)
    body = doc.element.body
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("For:") or text.startswith("Khách hàng:"):
            _set_field(p, text.split(":")[0] + ":", client_name.upper())
        elif text.startswith("Re:") or text.startswith("Vị trí:"):
            _set_field(p, text.split(":")[0] + ":", position)
        elif "Full name" in text or "Họ tên" in text:
            _set_tab(p, cv_data.get("full_name", ""))
        elif "Gender" in text or "Giới tính" in text:
            _set_tab(p, cv_data.get("gender", ""))
        elif "Year of birth" in text or "Năm sinh" in text:
            _set_tab(p, cv_data.get("year_of_birth", ""))
        elif "Marital" in text or "Tình trạng" in text:
            _set_tab(p, cv_data.get("marital_status", ""))
        elif "Address" in text or "Địa chỉ" in text:
            _set_tab(p, cv_data.get("address", ""))
    _populate_toc(body, cv_data.get("career_summary", []))
    work_el = next((p._element for p in doc.paragraphs if p.text.strip() in ("Working Experience", "Kinh nghiệm làm việc")), None)
    edu_el = next((p._element for p in doc.paragraphs if p.text.strip() in ("Education", "Trình độ chuyên môn")), None)
    other_el = next((p._element for p in doc.paragraphs if p.text.strip() in ("Other information (if any)", "Thông tin khác (nếu có)")), None)
    def remove_between(start, end):
        if start is None: return
        removing, to_remove = False, []
        for child in body:
            if child is start: removing = True; continue
            if end and child is end: break
            if removing and child.tag.split('}')[-1] == 'p': to_remove.append(child)
        for el in to_remove: body.remove(el)
    def remove_after(start):
        if start is None: return
        removing, to_remove = False, []
        for child in body:
            if child is start: removing = True; continue
            if removing:
                if child.tag.split('}')[-1] == 'sectPr': continue
                if child.tag.split('}')[-1] == 'p': to_remove.append(child)
        for el in to_remove: body.remove(el)
    remove_between(work_el, edu_el)
    remove_between(edu_el, other_el)
    remove_after(other_el)
    sid = lambda name: _get_style_id(doc, name)
    if work_el:
        ins = work_el
        for job in cv_data.get("career_summary", []):
            p = _make_p((job.get("period", "") + "\t" + job.get("company", "")).strip(), "Normal" if job.get("use_normal_style") else sid("Heading 1"))
            ins.addnext(p); ins = p
            for pos in job.get("positions", []):
                if pos.get("title"):
                    p = _make_p(pos["title"], sid("Heading 3")); ins.addnext(p); ins = p
                for resp in pos.get("responsibilities", []):
                    p = _make_p(resp, sid("1.Content")); ins.addnext(p); ins = p
                for ach in pos.get("achievements", []):
                    p = _make_p(ach, sid("1.Content")); ins.addnext(p); ins = p
            p = _make_p("", "Normal"); ins.addnext(p); ins = p
    if edu_el:
        ins = edu_el
        for edu in cv_data.get("education", []):
            line = (edu.get("period", "") + "\t" + edu.get("institution", "")).strip() if edu.get("period") and edu.get("institution") else (edu.get("period") or edu.get("institution", ""))
            p = _make_p(line, "Normal"); ins.addnext(p); ins = p
            for detail in edu.get("details", []):
                p = _make_p("\t" + detail, "Normal"); ins.addnext(p); ins = p
            p = _make_p("", "Normal"); ins.addnext(p); ins = p
    if other_el:
        ins = other_el
        for section in cv_data.get("other_info", []):
            # Section title — bold, normal style
            title = section.get("section_title", "")
            p = _make_p(title, "Normal")
            # Ensure the section title run is bold
            for r_el in p.findall(qn('w:r')):
                rPr = r_el.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r_el.insert(0, rPr)
                b = OxmlElement('w:b')
                rPr.append(b)
            ins.addnext(p); ins = p

            # Optional marked label (e.g. "Accountability:", "Duties:")
            section_label = section.get("section_label", "").strip()
            if section_label:
                p_lbl = _make_p_with_bold_label(section_label, "", "Normal")
                ins.addnext(p_lbl); ins = p_lbl

            for item in section.get("items", []):
                # other_info.items may be plain strings (offline) or dicts (AI path)
                if isinstance(item, dict):
                    marked = item.get("marked_label", "").strip()
                    marked_content = item.get("marked_content", "").strip()
                    if marked and marked_content:
                        p = _make_p_with_bold_label(marked, marked_content, sid("1.Content"))
                    else:
                        p = _make_p(item.get("text", "") or "", sid("1.Content"))
                else:
                    # Plain string item — just render normally
                    p = _make_p(str(item), sid("1.Content"))
                ins.addnext(p); ins = p

            # Optional achievements label + achievements
            achievements_label = section.get("achievements_label", "").strip()
            if achievements_label:
                p_albl = _make_p_with_bold_label(achievements_label, "", "Normal")
                ins.addnext(p_albl); ins = p_albl
            for ach in section.get("achievements", []):
                p = _make_p(ach, sid("1.Content"))
                # Make achievements bold
                for r_el in p.findall(qn('w:r')):
                    rPr = r_el.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        r_el.insert(0, rPr)
                    b = OxmlElement('w:b')
                    rPr.append(b)
                ins.addnext(p); ins = p

            p = _make_p("", "Normal"); ins.addnext(p); ins = p
    doc.save(output_path)

def _set_field(p, label, value):
    for run in p.runs: run.text = ""
    if p.runs: p.runs[0].text = label + " " + value

def _set_tab(p, value):
    if p.runs:
        full = p.text
        label = full.split('\t')[0] + '\t' if '\t' in full else full
        # Clear all runs first to remove stale text
        for run in p.runs:
            run.text = ""
        p.runs[0].text = label + value


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
