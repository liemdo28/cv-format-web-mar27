import json
import os
import re
from difflib import SequenceMatcher
from typing import Any, Dict, List, Set

from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RULES_PATH = os.path.join(BASE_DIR, "learning_store.json")

DEFAULT_RULES: Dict[str, Any] = {
    "aliases": {
        "full_name": ["name", "candidate_name", "full name", "họ tên", "ứng viên"],
        "email": ["e-mail", "mail"],
        "phone": ["mobile", "phone_number", "điện thoại", "số điện thoại"],
        "current_company": ["company", "company_name", "công ty", "công ty hiện tại"],
        "current_position": ["position", "job_title", "chức danh", "vị trí hiện tại"],
        "summary": ["profile", "professional_summary", "tóm tắt"],
        "experience": ["work_experience", "kinh nghiệm", "employment_history"],
        "education": ["học vấn", "academic"],
        "skills": ["kỹ năng", "core_skills"],
        "languages": ["ngôn ngữ", "language"],
        "address": ["location", "địa chỉ"],
    },
    "template_map": {
        "full_name": "full_name",
        "name": "full_name",
        "candidate_name": "full_name",
        "email": "email",
        "phone": "phone",
        "current_company": "current_company",
        "current_position": "current_position",
        "summary": "summary",
        "experience": "experience",
        "education": "education",
        "skills": "skills",
        "languages": "languages",
        "address": "address",
    },
}


def _normalize(text: str) -> str:
    return re.sub(r"[^a-z0-9_]+", "_", (text or "").strip().lower()).strip("_")


def load_rules() -> Dict[str, Any]:
    if not os.path.exists(RULES_PATH):
        with open(RULES_PATH, "w", encoding="utf-8") as f:
            json.dump(DEFAULT_RULES, f, ensure_ascii=False, indent=2)
        return DEFAULT_RULES
    with open(RULES_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def save_rules(rules: Dict[str, Any]):
    with open(RULES_PATH, "w", encoding="utf-8") as f:
        json.dump(rules, f, ensure_ascii=False, indent=2)


def _extract_email(text: str) -> str:
    m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    return m.group(0) if m else ""


def _extract_phone(text: str) -> str:
    m = re.search(r"(?:\+?\d[\d\s\-().]{7,}\d)", text)
    return m.group(0).strip() if m else ""


def _first_nontrivial_line(lines: List[str]) -> str:
    for line in lines:
        if len(line) > 2 and not re.search(r"@|\+?\d", line):
            return line
    return lines[0] if lines else ""


def _extract_section_blob(full_text: str, heading_patterns: List[str], stop_patterns: List[str]) -> str:
    heading_regex = r"(?:" + "|".join(heading_patterns) + r")"
    stop_regex = r"(?:" + "|".join(stop_patterns) + r")"
    pattern = rf"{heading_regex}\s*:?\s*(.+?)(?=\n\s*{stop_regex}\s*:|\Z)"
    m = re.search(pattern, full_text, flags=re.IGNORECASE | re.DOTALL)
    return (m.group(1).strip() if m else "")[:4000]


def extract_offline(cv_text: str) -> Dict[str, Any]:
    lines = [ln.strip() for ln in (cv_text or "").splitlines() if ln.strip()]
    full_text = "\n".join(lines)

    full_name = _first_nontrivial_line(lines)
    email = _extract_email(full_text)
    phone = _extract_phone(full_text)

    summary = _extract_section_blob(
        full_text,
        [r"summary", r"profile", r"professional summary", r"t[oó]m t[ắa]t"],
        [r"experience", r"education", r"skills", r"languages", r"projects?"],
    )
    experience_blob = _extract_section_blob(
        full_text,
        [r"experience", r"work experience", r"employment", r"kinh nghi[ệe]m"],
        [r"education", r"skills", r"languages", r"projects?"],
    )
    education_blob = _extract_section_blob(
        full_text,
        [r"education", r"academic", r"h[ọo]c v[ấa]n"],
        [r"skills", r"languages", r"projects?", r"certifications?"],
    )
    skills_blob = _extract_section_blob(
        full_text,
        [r"skills", r"core skills", r"k[ỹy] n[ăa]ng"],
        [r"languages", r"projects?", r"certifications?"],
    )
    languages_blob = _extract_section_blob(
        full_text,
        [r"languages?", r"ng[oô]n ng[ữu]"],
        [r"projects?", r"certifications?", r"references?"],
    )

    current_position = ""
    current_company = ""
    for line in lines[:140]:
        if " - " in line and len(line) <= 120:
            a, b = [x.strip() for x in line.split(" - ", 1)]
            if len(a) > 1 and len(b) > 1:
                current_position, current_company = a, b
                break

    return {
        "full_name": full_name,
        "gender": "",
        "year_of_birth": "",
        "marital_status": "",
        "address": "",
        "current_company": current_company,
        "current_position": current_position,
        "email": email,
        "phone": phone,
        "summary": summary,
        "career_summary": [
            {
                "period": "",
                "company": current_company.upper() if current_company else "",
                "company_description": "",
                "positions": [
                    {
                        "period": "",
                        "title": current_position,
                        "report_to": "",
                        "section_label": "",
                        "responsibilities": [x.strip() for x in experience_blob.split("\n") if x.strip()][:12],
                        "achievements_label": "",
                        "achievements": [],
                    }
                ] if current_position or experience_blob else [],
            }
        ] if (current_company or current_position or experience_blob) else [],
        "education": [
            {
                "period": "",
                "institution": "",
                "details": [x.strip() for x in education_blob.split("\n") if x.strip()][:8],
            }
        ] if education_blob else [],
        "other_info": [
            {"section_title": "SKILLS", "items": [x.strip() for x in skills_blob.split("\n") if x.strip()][:12]},
            {"section_title": "LANGUAGES", "items": [x.strip() for x in languages_blob.split("\n") if x.strip()][:8]},
            {"section_title": "SUMMARY", "items": [summary] if summary else []},
        ],
    }


def build_suggested_name_offline(cv_data: Dict[str, Any]) -> str:
    company = (cv_data.get("current_company") or "").strip()
    position = (cv_data.get("current_position") or "").strip()
    full_name = (cv_data.get("full_name") or "").strip()

    if not company:
        career = cv_data.get("career_summary") or []
        if career:
            company = (career[0].get("company") or "").strip().title()
    if not position:
        career = cv_data.get("career_summary") or []
        if career:
            positions = career[0].get("positions") or []
            if positions:
                position = (positions[0].get("title") or "").strip()

    parts = [p for p in [company, position, full_name] if p]
    return " - ".join(parts)


def _collect_placeholders(doc: Document) -> Set[str]:
    placeholders: Set[str] = set()

    def grab(text: str):
        for token in re.findall(r"\{\{([^{}]+)\}\}", text or ""):
            placeholders.add(token.strip())

    for p in doc.paragraphs:
        grab(p.text)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    grab(p.text)

    for section in doc.sections:
        for p in section.header.paragraphs:
            grab(p.text)
        for p in section.footer.paragraphs:
            grab(p.text)

    return placeholders


def _best_field_for_placeholder(placeholder: str, rules: Dict[str, Any]) -> tuple[str, float]:
    p = _normalize(placeholder)
    template_map = rules.get("template_map", {})
    aliases = rules.get("aliases", {})

    if p in template_map:
        return template_map[p], 1.0

    best_field = ""
    best_score = 0.0
    for canonical, alias_list in aliases.items():
        candidates = [canonical, *alias_list]
        for c in candidates:
            score = SequenceMatcher(None, p, _normalize(c)).ratio()
            if score > best_score:
                best_score = score
                best_field = canonical

    if best_score < 0.58:
        return "", best_score
    return best_field, best_score


def _replace_token_everywhere(doc: Document, token: str, value: str):
    def replace_in_paragraph(paragraph):
        if token not in paragraph.text:
            return
        for run in paragraph.runs:
            if token in run.text:
                run.text = run.text.replace(token, value)

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p)
        for p in section.footer.paragraphs:
            replace_in_paragraph(p)


def fill_template_offline(template_path: str, output_path: str, cv_data: Dict[str, Any]) -> Dict[str, Any]:
    rules = load_rules()
    doc = Document(template_path)
    placeholders = sorted(_collect_placeholders(doc))

    mapping_results = []

    for ph in placeholders:
        token = "{{" + ph + "}}"
        canonical, confidence = _best_field_for_placeholder(ph, rules)

        value = ""
        if canonical:
            raw = cv_data.get(canonical, "")
            if isinstance(raw, list):
                flattened = []
                for item in raw:
                    if isinstance(item, str):
                        flattened.append(item)
                    elif isinstance(item, dict):
                        flattened.extend([str(v) for v in item.values() if isinstance(v, str) and v.strip()])
                value = "\n".join(flattened[:8])
            elif isinstance(raw, dict):
                value = "\n".join([str(v) for v in raw.values() if isinstance(v, str) and v.strip()][:8])
            else:
                value = str(raw or "")

        _replace_token_everywhere(doc, token, value)
        mapping_results.append({
            "placeholder": ph,
            "mappedField": canonical,
            "confidence": round(confidence, 3),
            "filled": bool(value.strip()),
        })

    doc.save(output_path)

    review_required = [m for m in mapping_results if (not m["mappedField"] or m["confidence"] < 0.7)]
    return {
        "mappings": mapping_results,
        "reviewRequired": review_required,
    }


def learn_mapping(placeholder: str, canonical_field: str):
    rules = load_rules()
    p = _normalize(placeholder)
    c = _normalize(canonical_field)

    rules.setdefault("template_map", {})[p] = c
    rules.setdefault("aliases", {}).setdefault(c, [])
    if p not in rules["aliases"][c]:
        rules["aliases"][c].append(p)

    save_rules(rules)


def _extract_label_candidates(text: str) -> List[str]:
    labels: List[str] = []
    for ln in (text or "").splitlines():
        line = ln.strip()
        if not line or len(line) > 60:
            continue
        if ":" in line:
            label = line.split(":", 1)[0].strip()
            if 1 < len(label) <= 40:
                labels.append(label)
    return labels


def _canonical_from_label(label: str) -> tuple[str, float]:
    candidates = {
        "full_name": ["full name", "họ tên", "candidate name", "name"],
        "email": ["email", "e-mail", "mail"],
        "phone": ["phone", "mobile", "điện thoại", "số điện thoại"],
        "address": ["address", "địa chỉ", "location"],
        "experience": ["working experience", "work experience", "experience", "kinh nghiệm"],
        "education": ["education", "trình độ chuyên môn", "học vấn", "academic"],
        "skills": ["skills", "kỹ năng"],
        "languages": ["languages", "ngôn ngữ"],
        "summary": ["summary", "profile", "tóm tắt"],
    }

    norm = _normalize(label)
    best_key = ""
    best_score = 0.0
    for key, alias_list in candidates.items():
        for alias in alias_list:
            score = SequenceMatcher(None, norm, _normalize(alias)).ratio()
            if score > best_score:
                best_score = score
                best_key = key
    return best_key, best_score


def learn_from_training_pair(raw_text: str, done_text: str) -> Dict[str, Any]:
    """
    Learn additional format aliases from two documents:
    - raw_text: unfinished CV
    - done_text: formatted/final CV
    """
    rules = load_rules()

    learned_aliases: Dict[str, List[str]] = {}
    learned_mappings: Dict[str, str] = {}

    # 1) Learn from label-style lines in completed file (e.g., "Full name:")
    label_candidates = _extract_label_candidates(done_text)
    for label in label_candidates:
        canonical, score = _canonical_from_label(label)
        if not canonical or score < 0.64:
            continue

        nlabel = _normalize(label)
        rules.setdefault("aliases", {}).setdefault(canonical, [])
        if nlabel and nlabel not in rules["aliases"][canonical]:
            rules["aliases"][canonical].append(nlabel)
            learned_aliases.setdefault(canonical, []).append(nlabel)

        rules.setdefault("template_map", {})[nlabel] = canonical
        learned_mappings[nlabel] = canonical

    # 2) Learn uppercase section titles from completed file
    section_lines = []
    for ln in (done_text or "").splitlines():
        line = ln.strip()
        if 2 < len(line) <= 40 and re.match(r"^[A-ZÀ-Ỹ0-9\s&/().,-]+$", line):
            section_lines.append(line)

    for sec in section_lines:
        canonical, score = _canonical_from_label(sec)
        if not canonical or score < 0.60:
            continue
        nsec = _normalize(sec)
        rules.setdefault("aliases", {}).setdefault(canonical, [])
        if nsec and nsec not in rules["aliases"][canonical]:
            rules["aliases"][canonical].append(nsec)
            learned_aliases.setdefault(canonical, []).append(nsec)

    # 3) Save a compact training memory snapshot
    rules.setdefault("training_pairs", [])
    rules["training_pairs"].append({
        "raw_preview": (raw_text or "")[:200],
        "done_preview": (done_text or "")[:200],
        "learned_count": sum(len(v) for v in learned_aliases.values()),
    })
    rules["training_pairs"] = rules["training_pairs"][-20:]

    save_rules(rules)

    return {
        "learnedAliases": learned_aliases,
        "learnedMappings": learned_mappings,
        "labelCandidates": len(label_candidates),
        "sectionCandidates": len(section_lines),
    }


def load_learning_store() -> Dict[str, Any]:
    return load_rules()


def reset_learning_store() -> Dict[str, Any]:
    save_rules(DEFAULT_RULES)
    return DEFAULT_RULES
